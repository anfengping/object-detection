# -*- coding: utf-8 -*-
from pathlib import Path
import os, math, zipfile, shutil, subprocess, textwrap, datetime
from xml.sax.saxutils import escape
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
import svgwrite
import cairosvg

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "output"
FIG_SVG = OUT / "Figures_SVG"
FIG_PNG = OUT / "Figures_PNG"
FIG_VSDX = OUT / "Visio_Editable_Figures"
for p in [OUT, FIG_SVG, FIG_PNG, FIG_VSDX]: p.mkdir(parents=True, exist_ok=True)

TITLE = "输配电线路无人机巡检与树障隐患智能感知关键技术研究与应用"

# ----------------------- DOCX utilities -----------------------
def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_text(cell, text, bold=False, size=9.5, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(text)); r.bold = bold; r.font.size = Pt(size)
    r.font.name = 'Microsoft YaHei'; r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if color: r.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr(); tblHeader = OxmlElement('w:tblHeader'); tblHeader.set(qn('w:val'), 'true'); trPr.append(tblHeader)

def set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m, v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node = tcMar.find(qn('w:'+m))
        if node is None: node = OxmlElement('w:'+m); tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'),'dxa')

def add_table(doc, headers, rows, widths=None, font_size=8.8):
    table = doc.add_table(rows=1, cols=len(headers)); table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.style='Table Grid'
    hdr = table.rows[0]; set_repeat_table_header(hdr)
    for i,h in enumerate(headers):
        set_cell_text(hdr.cells[i],h,True,9.2,(255,255,255)); set_cell_shading(hdr.cells[i],'315B7D'); set_cell_margins(hdr.cells[i])
    for ri,row in enumerate(rows):
        cells = table.add_row().cells
        for i,val in enumerate(row):
            set_cell_text(cells[i],val,False,font_size)
            if ri%2==1: set_cell_shading(cells[i],'EAF2F8')
            set_cell_margins(cells[i])
    if widths:
        for row in table.rows:
            for i,w in enumerate(widths): row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table

def set_run_font(run, name='SimSun', size=10.5, bold=None, color=None):
    run.font.name=name; run._element.rPr.rFonts.set(qn('w:eastAsia'),name); run.font.size=Pt(size)
    if bold is not None: run.bold=bold
    if color: run.font.color.rgb=RGBColor(*color)

def add_p(doc, text='', style=None, bold_prefix=None, align=None, indent=True, spacing=1.35, before=0, after=3):
    p=doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r=p.add_run(bold_prefix); set_run_font(r,'Microsoft YaHei',10.5,True)
        r=p.add_run(text[len(bold_prefix):]); set_run_font(r)
    else:
        r=p.add_run(text); set_run_font(r)
    if align is not None: p.alignment=align
    pf=p.paragraph_format; pf.line_spacing=spacing; pf.space_before=Pt(before); pf.space_after=Pt(after)
    if indent and style not in ['Title','Subtitle','Heading 1','Heading 2','Heading 3','Heading 4']:
        pf.first_line_indent=Cm(0.74)
    return p

def add_bullet(doc, text):
    p=doc.add_paragraph(style='List Bullet'); r=p.add_run(text); set_run_font(r); p.paragraph_format.line_spacing=1.25; p.paragraph_format.space_after=Pt(2); return p

def add_equation(doc, eq, num, variables):
    table=doc.add_table(rows=1, cols=3); table.alignment=WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width=Cm(1.2); table.columns[1].width=Cm(14.5); table.columns[2].width=Cm(1.2)
    p=table.cell(0,1).paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    omath_para=OxmlElement('m:oMathPara'); omath=OxmlElement('m:oMath'); mr=OxmlElement('m:r')
    mrpr=OxmlElement('m:rPr'); sty=OxmlElement('m:sty'); sty.set(qn('m:val'),'p'); mrpr.append(sty); mr.append(mrpr)
    mt=OxmlElement('m:t'); mt.text=eq; mr.append(mt); omath.append(mr); omath_para.append(omath); p._p.append(omath_para)
    pn=table.cell(0,2).paragraphs[0]; pn.alignment=WD_ALIGN_PARAGRAPH.RIGHT; rn=pn.add_run(f"（{num}）"); set_run_font(rn,'SimSun',10.5)
    for c in table.rows[0].cells:
        tcPr=c._tc.get_or_add_tcPr(); tcBorders=tcPr.first_child_found_in('w:tcBorders')
        if tcBorders is None: tcBorders=OxmlElement('w:tcBorders'); tcPr.append(tcBorders)
        for edge in ['top','left','bottom','right','insideH','insideV']:
            e=OxmlElement('w:'+edge); e.set(qn('w:val'),'nil'); tcBorders.append(e)
    add_p(doc, f"式中，{variables}", indent=True, spacing=1.25, after=4)

def add_figure(doc, idx, title, png_path, vsdx_name, width=6.5):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(3)
    p.add_run().add_picture(str(png_path), width=Inches(width))
    cp=doc.add_paragraph(); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=cp.add_run(f"图 {idx}  {title}"); set_run_font(r,'Microsoft YaHei',10,True)
    sp=doc.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=sp.add_run(f"可编辑源文件：{vsdx_name}"); set_run_font(r,'SimSun',8.5,False,(90,90,90)); sp.paragraph_format.space_after=Pt(5)

def add_heading(doc, text, level=1):
    p=doc.add_paragraph(style=f'Heading {level}'); r=p.add_run(text); set_run_font(r,'Microsoft YaHei',16 if level==1 else 13 if level==2 else 11,True); return p

# ----------------------- Editable Visio package -----------------------
COLORS=['D9EAF7','E2F0D9','FFF2CC','FCE4D6','E4DFEC','DDEBF7','EDEDED']

def wrap_label(s, width=12):
    lines=[]
    for part in str(s).split('\n'):
        if len(part)<=width: lines.append(part)
        else:
            for i in range(0,len(part),width): lines.append(part[i:i+width])
    return '\n'.join(lines)

def make_diagram(idx, title, nodes, edges, subtitle=''):
    W,H=1400,820
    svg=svgwrite.Drawing(str(FIG_SVG/f"Figure_{idx:02d}.svg"), size=(W,H), viewBox=f"0 0 {W} {H}")
    svg.add(svg.rect(insert=(0,0),size=(W,H),fill='white'))
    marker=svg.marker(insert=(10,5),size=(10,10),orient='auto',id='arrow'); marker.add(svg.path(d='M0,0 L10,5 L0,10 z',fill='#4F81BD')); svg.defs.add(marker)
    svg.add(svg.text(title,insert=(W/2,55),text_anchor='middle',font_size=30,font_family='Microsoft YaHei',font_weight='bold',fill='#17365D'))
    if subtitle: svg.add(svg.text(subtitle,insert=(W/2,86),text_anchor='middle',font_size=17,font_family='Microsoft YaHei',fill='#666666'))
    byid={n['id']:n for n in nodes}
    for a,b,label in edges:
        na,nb=byid[a],byid[b]; x1=na['x']+na['w']/2; y1=na['y']+na['h']/2; x2=nb['x']+nb['w']/2; y2=nb['y']+nb['h']/2
        svg.add(svg.line(start=(x1,y1),end=(x2,y2),stroke='#4F81BD',stroke_width=2.4,marker_end='url(#arrow)'))
        if label: svg.add(svg.text(label,insert=((x1+x2)/2,(y1+y2)/2-6),text_anchor='middle',font_size=14,font_family='Microsoft YaHei',fill='#444'))
    for i,n in enumerate(nodes):
        fill='#'+n.get('fill',COLORS[i%len(COLORS)])
        svg.add(svg.rect(insert=(n['x'],n['y']),size=(n['w'],n['h']),rx=10,ry=10,fill=fill,stroke='#507AA3',stroke_width=2))
        lines=wrap_label(n['text'],n.get('wrap',12)).split('\n'); line_h=22; start_y=n['y']+n['h']/2-(len(lines)-1)*line_h/2
        for j,line in enumerate(lines): svg.add(svg.text(line,insert=(n['x']+n['w']/2,start_y+j*line_h+5),text_anchor='middle',font_size=n.get('fs',18),font_family='Microsoft YaHei',font_weight='bold' if n.get('bold',False) else 'normal',fill='#1F1F1F'))
    svg.save()
    cairosvg.svg2png(url=str(FIG_SVG/f"Figure_{idx:02d}.svg"),write_to=str(FIG_PNG/f"Figure_{idx:02d}.png"),output_width=1400,output_height=820)
    build_vsdx(FIG_VSDX/f"Figure_{idx:02d}_{title.replace('—','_').replace('/','_')}.vsdx",title,nodes,edges,W,H)
    return FIG_PNG/f"Figure_{idx:02d}.png", next(FIG_VSDX.glob(f"Figure_{idx:02d}_*.vsdx")).name

def build_vsdx(path,title,nodes,edges,W,H):
    ns='http://schemas.microsoft.com/office/visio/2012/main'; rel='http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    page_w,page_h=13.333,7.619
    def vx(x): return x/W*page_w
    def vy(y): return page_h-y/H*page_h
    shapes=[]; sid=1; idmap={}
    # title
    shapes.append(f'''<Shape ID="{sid}" NameU="Title" Type="Shape"><Cell N="PinX" V="{page_w/2:.4f}"/><Cell N="PinY" V="{page_h-0.35:.4f}"/><Cell N="Width" V="12.0"/><Cell N="Height" V="0.45"/><Cell N="LocPinX" V="6"/><Cell N="LocPinY" V="0.225"/><Cell N="LinePattern" V="0"/><Cell N="FillPattern" V="0"/><Section N="Character"><Row IX="0"><Cell N="Font" V="0"/><Cell N="Size" V="0.22"/><Cell N="Style" V="1"/><Cell N="Color" V="#17365D"/></Row></Section><Section N="Paragraph"><Row IX="0"><Cell N="HorzAlign" V="1"/></Row></Section><Text>{escape(title)}</Text></Shape>'''); sid+=1
    for i,n in enumerate(nodes):
        idmap[n['id']]=sid; x=vx(n['x']+n['w']/2); y=vy(n['y']+n['h']/2); w=vx(n['w']); h=n['h']/H*page_h; fill='#'+n.get('fill',COLORS[i%len(COLORS)])
        txt=escape(wrap_label(n['text'],n.get('wrap',12)))
        shapes.append(f'''<Shape ID="{sid}" NameU="Box.{sid}" Type="Shape"><Cell N="PinX" V="{x:.4f}"/><Cell N="PinY" V="{y:.4f}"/><Cell N="Width" V="{w:.4f}"/><Cell N="Height" V="{h:.4f}"/><Cell N="LocPinX" V="{w/2:.4f}"/><Cell N="LocPinY" V="{h/2:.4f}"/><Cell N="FillForegnd" V="{fill}"/><Cell N="LineColor" V="#507AA3"/><Cell N="LineWeight" V="0.018"/><Section N="Geometry" IX="0"><Row T="MoveTo" IX="1"><Cell N="X" V="0"/><Cell N="Y" V="0"/></Row><Row T="LineTo" IX="2"><Cell N="X" V="{w:.4f}"/><Cell N="Y" V="0"/></Row><Row T="LineTo" IX="3"><Cell N="X" V="{w:.4f}"/><Cell N="Y" V="{h:.4f}"/></Row><Row T="LineTo" IX="4"><Cell N="X" V="0"/><Cell N="Y" V="{h:.4f}"/></Row><Row T="LineTo" IX="5"><Cell N="X" V="0"/><Cell N="Y" V="0"/></Row></Section><Section N="Character"><Row IX="0"><Cell N="Font" V="0"/><Cell N="Size" V="0.14"/><Cell N="Color" V="#1F1F1F"/></Row></Section><Section N="Paragraph"><Row IX="0"><Cell N="HorzAlign" V="1"/></Row></Section><Text>{txt}</Text></Shape>'''); sid+=1
    connects=[]
    for a,b,label in edges:
        na,nb=next(n for n in nodes if n['id']==a),next(n for n in nodes if n['id']==b)
        x1,y1=vx(na['x']+na['w']/2),vy(na['y']+na['h']/2); x2,y2=vx(nb['x']+nb['w']/2),vy(nb['y']+nb['h']/2)
        w=abs(x2-x1) or 0.01; h=abs(y2-y1) or 0.01; px=(x1+x2)/2; py=(y1+y2)/2
        bx=x1-px+w/2; by=y1-py+h/2; ex=x2-px+w/2; ey=y2-py+h/2
        shapes.append(f'''<Shape ID="{sid}" NameU="Connector.{sid}" Type="Shape"><Cell N="PinX" V="{px:.4f}"/><Cell N="PinY" V="{py:.4f}"/><Cell N="Width" V="{w:.4f}"/><Cell N="Height" V="{h:.4f}"/><Cell N="LocPinX" V="{w/2:.4f}"/><Cell N="LocPinY" V="{h/2:.4f}"/><Cell N="FillPattern" V="0"/><Cell N="LineColor" V="#4F81BD"/><Cell N="LineWeight" V="0.018"/><Cell N="EndArrow" V="13"/><Section N="Geometry" IX="0"><Row T="MoveTo" IX="1"><Cell N="X" V="{bx:.4f}"/><Cell N="Y" V="{by:.4f}"/></Row><Row T="LineTo" IX="2"><Cell N="X" V="{ex:.4f}"/><Cell N="Y" V="{ey:.4f}"/></Row></Section></Shape>''')
        sid+=1
    page=f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?><PageContents xmlns="{ns}" xmlns:r="{rel}"><PageSheet><Cell N="PageWidth" V="{page_w}"/><Cell N="PageHeight" V="{page_h}"/><Cell N="DrawingScale" V="1"/><Cell N="PageScale" V="1"/></PageSheet><Shapes>{''.join(shapes)}</Shapes></PageContents>'''
    parts={
    '[Content_Types].xml':'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"><Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/><Default Extension="xml" ContentType="application/xml"/><Override PartName="/visio/document.xml" ContentType="application/vnd.ms-visio.drawing.main+xml"/><Override PartName="/visio/pages/pages.xml" ContentType="application/vnd.ms-visio.pages+xml"/><Override PartName="/visio/pages/page1.xml" ContentType="application/vnd.ms-visio.page+xml"/><Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/><Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/></Types>''',
    '_rels/.rels':'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.microsoft.com/visio/2010/relationships/document" Target="visio/document.xml"/><Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/><Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/></Relationships>''',
    'visio/document.xml':f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?><VisioDocument xmlns="{ns}" xmlns:r="{rel}"><DocumentSettings/><Pages><Rel r:id="rId1"/></Pages></VisioDocument>''',
    'visio/_rels/document.xml.rels':'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.microsoft.com/visio/2010/relationships/pages" Target="pages/pages.xml"/></Relationships>''',
    'visio/pages/pages.xml':f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Pages xmlns="{ns}" xmlns:r="{rel}"><Page ID="0" NameU="Page-1" Name="Page-1"><Rel r:id="rId1"/></Page></Pages>''',
    'visio/pages/_rels/pages.xml.rels':'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.microsoft.com/visio/2010/relationships/page" Target="page1.xml"/></Relationships>''',
    'visio/pages/page1.xml':page,
    'docProps/core.xml':f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?><cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance"><dc:title>{escape(title)}</dc:title><dc:creator>OpenAI temporary build</dc:creator><dcterms:created xsi:type="dcterms:W3CDTF">2026-07-16T00:00:00Z</dcterms:created></cp:coreProperties>''',
    'docProps/app.xml':'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"><Application>Microsoft Visio</Application><AppVersion>16.0000</AppVersion></Properties>'''}
    with zipfile.ZipFile(path,'w',zipfile.ZIP_DEFLATED) as z:
        for k,v in parts.items(): z.writestr(k,v)

# ----------------------- Figure definitions -----------------------
fig_defs=[]
def N(i,t,x,y,w=210,h=86,fill=None,wrap=11,bold=False,fs=18): return {'id':i,'text':t,'x':x,'y':y,'w':w,'h':h,'fill':fill or COLORS[(int(i[-1]) if i[-1].isdigit() else 0)%len(COLORS)],'wrap':wrap,'bold':bold,'fs':fs}

fig_defs.append((1,'项目总体技术体系与创新链',[
N('a','国家战略与行业需求',70,120,230,90,'D9EAF7'),N('b','空中多源感知层',350,120,230,90,'E2F0D9'),N('c','智能认知与决策层',630,120,230,90,'FFF2CC'),N('d','端—边—云平台层',910,120,230,90,'FCE4D6'),N('e','业务闭环与效益层',1170,120,190,90,'E4DFEC'),
N('f','自主定位\n导线跟随\n安全避障',250,360,230,125,'DDEBF7'),N('g','部件识别\n缺陷诊断\n树障量测',585,360,230,125,'E2F0D9'),N('h','告警工单\n模型迭代\n证据留痕',920,360,230,125,'FFF2CC'),N('i','产品化\n规模应用\n综合效益',1150,590,190,100,'FCE4D6')],[('a','b','需求牵引'),('b','c','数据流'),('c','d','结果流'),('d','e','业务流'),('b','f',''),('c','g',''),('d','h',''),('h','i','转化')],'形成“感知—认知—决策—执行—反馈—转化”一体化闭环'))
fig_defs.append((2,'关键科学问题—技术突破—成果形态映射',[
N('a','多源异构数据\n时空不一致',80,130,240,100),N('b','小目标缺陷\n少样本与域偏移',80,350,240,100),N('c','树障风险\n空间性与动态性',80,570,240,100),
N('d','时空标定与\n鲁棒融合定位',540,130,260,100,'E2F0D9'),N('e','级联注意力+\n自监督+主动学习',540,350,260,100,'FFF2CC'),N('f','语义点云融合+\n数字孪生风险预测',540,570,260,100,'FCE4D6'),
N('g','自主巡检算法\n嵌入式原型',1030,130,260,100,'D9EAF7'),N('h','识别模型\n数据集与工具链',1030,350,260,100,'DDEBF7'),N('i','树障量测模块\n风险管理平台',1030,570,260,100,'E4DFEC')],[('a','d','科学问题'),('b','e','科学问题'),('c','f','科学问题'),('d','g','技术成果'),('e','h','技术成果'),('f','i','技术成果')],''))
fig_defs.append((3,'多传感器时空标定与鲁棒融合定位框架',[
N('a','GNSS/RTK',70,130),N('b','IMU',70,300),N('c','LiDAR',70,470),N('d','可见光/红外相机',70,640),
N('e','时间同步\n偏移估计',390,205,230,100,'FFF2CC'),N('f','坐标外参\n联合标定',390,485,230,100,'E2F0D9'),N('g','误差状态滤波/\n因子图优化',760,330,270,120,'D9EAF7'),N('h','连续位姿、速度\n协方差与健康度',1110,330,240,120,'FCE4D6')],[('a','e',''),('b','e',''),('c','f',''),('d','f',''),('e','g','时间约束'),('f','g','空间约束'),('g','h','融合输出')],''))
fig_defs.append((4,'感知质量约束的自主飞行与安全控制闭环',[
N('a','任务与电子围栏',80,140),N('b','三维环境地图',80,350),N('c','成像质量评价',80,560),N('d','多目标航迹优化',480,250,260,110,'FFF2CC'),N('e','概率安全约束',480,500,260,110,'FCE4D6'),N('f','模型预测控制',850,370,250,110,'D9EAF7'),N('g','无人机执行机构',1160,370,200,110,'E2F0D9'),N('h','在线反馈与重规划',850,650,250,90,'E4DFEC')],[('a','d','任务约束'),('b','d','环境约束'),('c','d','质量约束'),('b','e','障碍概率'),('d','f','参考轨迹'),('e','f','安全边界'),('f','g','控制量'),('g','h','状态反馈'),('h','d','滚动优化')],''))
fig_defs.append((5,'巡检数据治理、主动学习与持续训练闭环',[
N('a','现场影像/点云\n任务日志',70,180,230,100),N('b','质量筛查与\n时空配准',370,180,230,100,'E2F0D9'),N('c','自动预标注与\n人机复核',670,180,230,100,'FFF2CC'),N('d','版本化数据集\n场景谱系',970,180,230,100,'FCE4D6'),N('e','训练、验证\n轻量化部署',970,500,230,100,'D9EAF7'),N('f','低置信度/误检/\n漏检回流',670,500,230,100,'E4DFEC'),N('g','主动采样与\n困难样本生成',370,500,230,100,'DDEBF7'),N('h','线上监测与\n模型健康度',70,500,230,100,'E2F0D9')],[('a','b',''),('b','c',''),('c','d',''),('d','e','训练'),('e','h','发布'),('h','f','异常样本'),('f','g','筛选'),('g','c','增量标注')],''))
fig_defs.append((6,'级联注意力—自监督—域泛化识别网络',[
N('a','多尺度图像输入',60,340,210,90),N('b','卷积骨干\n局部纹理',330,220,220,100,'D9EAF7'),N('c','通道—空间—\n跨层级联注意',650,220,250,100,'FFF2CC'),N('d','自监督对比学习\n未标注影像预训练',330,500,250,100,'E2F0D9'),N('e','域泛化与\n风格扰动',650,500,250,100,'FCE4D6'),N('f','多尺度特征融合',980,340,230,100,'E4DFEC'),N('g','部件定位、分类\n置信度与不确定度',1210,340,170,120,'DDEBF7')],[('a','b',''),('a','d',''),('b','c',''),('d','e',''),('c','f',''),('e','f',''),('f','g','')],''))
fig_defs.append((7,'卷积增强Transformer缺陷检测架构',[
N('a','高分辨率巡检图像',60,340,210,90),N('b','Patch Embedding\n+局部卷积',330,340,220,100,'D9EAF7'),N('c','窗口注意力\n+LWC',630,180,220,100,'FFF2CC'),N('d','移位窗口注意力\n+GWC',630,500,220,100,'FCE4D6'),N('e','跨层多尺度金字塔',950,340,230,110,'E2F0D9'),N('f','轻量检测头',1210,240,170,90,'DDEBF7'),N('g','分割/关键点头',1210,500,170,90,'E4DFEC')],[('a','b',''),('b','c',''),('b','d',''),('c','e','局部—全局'),('d','e','跨窗口'),('e','f','缺陷框'),('e','g','几何语义')],''))
fig_defs.append((8,'图像语义—LiDAR点云跨模态融合流程',[
N('a','图像语义分割\n导线/树冠/杆塔',70,180,240,110),N('b','LiDAR点云\n去噪与运动补偿',70,500,240,110),N('c','相机—LiDAR\n外参与时间同步',430,340,260,110,'FFF2CC'),N('d','语义投影与\n点级标签传播',800,340,250,110,'E2F0D9'),N('e','导线点集拟合\n树冠聚类',1130,220,220,110,'D9EAF7'),N('f','净空距离与\n置信区间',1130,500,220,110,'FCE4D6')],[('a','c','语义'),('b','c','几何'),('c','d','对齐'),('d','e','结构提取'),('d','f','空间量测'),('e','f','几何约束')],''))
fig_defs.append((9,'树障数字孪生、动态预测与分级处置',[
N('a','档距三维数字孪生',70,160,250,100),N('b','导线悬链线/弧垂\n风偏模型',70,420,250,110),N('c','树冠几何、树种\n生长率与气象',70,650,250,100),N('d','时序状态估计',480,300,250,110,'FFF2CC'),N('e','未来净空概率分布',480,570,250,110,'E2F0D9'),N('f','多因素风险指数',850,420,240,110,'FCE4D6'),N('g','提示/一般/严重/紧急',1160,300,210,100,'D9EAF7'),N('h','复飞、修剪、停电\n处置与复核',1160,570,210,110,'E4DFEC')],[('a','d','空间基准'),('b','d','导线状态'),('c','e','生长预测'),('d','e','状态外推'),('e','f','概率风险'),('f','g','分级'),('f','h','决策')],''))
fig_defs.append((10,'端—边—云协同推理与MLOps闭环',[
N('a','无人机/固定终端',60,180,230,100),N('b','边缘质量筛查\n快速推理与缓存',360,180,250,110,'E2F0D9'),N('c','通信与任务卸载',690,180,230,100,'FFF2CC'),N('d','云端高精度推理\nGIS与时序分析',1000,180,260,110,'D9EAF7'),N('e','告警、工单与处置',1000,520,260,110,'FCE4D6'),N('f','模型注册、评测\n灰度发布与回滚',690,520,230,110,'E4DFEC'),N('g','数据回流与\n主动学习',360,520,250,110,'DDEBF7'),N('h','运维人员与资产台账',60,520,230,100,'E2F0D9')],[('a','b','采集'),('b','c','任务'),('c','d','卸载'),('d','e','结果'),('e','h','闭环'),('d','f','模型监控'),('f','g','更新'),('g','b','增量模型')],''))
fig_defs.append((11,'技术先进性对标与成熟度提升路径',[
N('a','单一人工/遥控巡检',80,180,250,90),N('b','单任务图像检测',80,370,250,90),N('c','多模态但离线分析',80,560,250,90),N('d','本项目：自主采集+\n多模态认知+闭环处置',520,350,300,130,'FFF2CC'),N('e','算法原型 TRL4—5',980,150,240,90,'D9EAF7'),N('f','系统样机 TRL6',980,350,240,90,'E2F0D9'),N('g','典型场景示范 TRL7',980,550,240,90,'FCE4D6'),N('h','标准化产品与规模推广',1220,350,170,110,'E4DFEC')],[('a','d','能力跃迁'),('b','d','能力跃迁'),('c','d','能力跃迁'),('d','e','研发'),('e','f','集成'),('f','g','验证'),('g','h','转化')],''))
fig_defs.append((12,'客观评价与申报证据链',[
N('a','原始试验数据\n日志与影像',70,180,230,100),N('b','算法/系统测试报告',360,180,230,100,'E2F0D9'),N('c','第三方检测\nCMA/CNAS',650,180,230,100,'FFF2CC'),N('d','验收、查新与\n成果评价',940,180,230,100,'FCE4D6'),N('e','专利、软著、论文\n标准与应用证明',940,520,230,110,'D9EAF7'),N('f','财务凭证与\n专项审计',650,520,230,100,'E4DFEC'),N('g','创新点—单位—\n完成人映射',360,520,230,100,'DDEBF7'),N('h','奖项申报事实库',70,520,230,100,'E2F0D9')],[('a','b','可复现'),('b','c','独立验证'),('c','d','水平结论'),('d','e','成果支撑'),('e','f','效益核验'),('f','g','贡献核定'),('g','h','归档'),('h','a','追溯')],''))
fig_defs.append((13,'产学研Stage-Gate协同治理与知识产权机制',[
N('a','需求冻结 G0',60,180,190,90),N('b','算法原型 G1',300,180,190,90,'E2F0D9'),N('c','系统联调 G2',540,180,190,90,'FFF2CC'),N('d','现场验证 G3',780,180,190,90,'FCE4D6'),N('e','第三方评价 G4',1020,180,190,90,'D9EAF7'),N('f','转化推广 G5',1230,180,150,90,'E4DFEC'),N('g','技术委员会',230,500,220,100,'DDEBF7'),N('h','应用示范组',590,500,220,100,'E2F0D9'),N('i','知识产权与\n成果转化组',950,500,240,100,'FFF2CC')],[('a','b',''),('b','c',''),('c','d',''),('d','e',''),('e','f',''),('g','c','接口/基线'),('h','d','场景/数据'),('i','e','权属/收益')],''))
fig_defs.append((14,'成果产品化、场景复制与综合效益传导',[
N('a','核心算法与数据资产',60,180,230,100),N('b','边缘终端/算法SDK',360,180,230,100,'E2F0D9'),N('c','在线平台与风险模块',660,180,230,100,'FFF2CC'),N('d','巡检服务与联合方案',960,180,230,100,'FCE4D6'),N('e','电网/矿山/园区/应急',1190,180,190,110,'D9EAF7'),N('f','效率提升与成本节约',960,520,230,100,'E4DFEC'),N('g','停电风险与安全暴露降低',660,520,230,110,'DDEBF7'),N('h','产业收入、人才与标准',360,520,230,100,'E2F0D9'),N('i','经济—社会—生态综合效益',60,520,230,110,'FFF2CC')],[('a','b','工程化'),('b','c','集成'),('c','d','服务化'),('d','e','应用'),('e','f','直接效益'),('f','g','风险效益'),('g','h','产业效益'),('h','i','综合价值')],''))
fig_defs.append((15,'项目开发历程与持续演进路线',[
N('a','需求分析与\n方案论证',60,330,200,100),N('b','核心算法攻关',300,330,200,100,'E2F0D9'),N('c','缺陷与树障\n融合感知',540,330,200,100,'FFF2CC'),N('d','平台集成与\n边缘部署',780,330,200,100,'FCE4D6'),N('e','现场试点与\n第三方评价',1020,330,200,100,'D9EAF7'),N('f','产品定型、标准\n与规模推广',1240,330,150,110,'E4DFEC'),N('g','数据—模型—应用持续反馈',430,600,540,90,'DDEBF7')],[('a','b','G0'),('b','c','G1'),('c','d','G2'),('d','e','G3'),('e','f','G4/G5'),('e','g','问题回流'),('g','b','持续迭代')],''))

fig_paths={}
for args in fig_defs:
    idx,title,nodes,edges,subtitle=args
    fig_paths[idx]=make_diagram(idx,title,nodes,edges,subtitle)

# ----------------------- Build document -----------------------
doc=Document()
sec=doc.sections[0]; sec.top_margin=Cm(2.3); sec.bottom_margin=Cm(2.2); sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.3)
styles=doc.styles
styles['Normal'].font.name='SimSun'; styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'SimSun'); styles['Normal'].font.size=Pt(10.5)
for lvl,size in [(1,16),(2,14),(3,12),(4,11)]:
    s=styles[f'Heading {lvl}']; s.font.name='Microsoft YaHei'; s._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft YaHei'); s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor(31,78,121)
    s.paragraph_format.space_before=Pt(8); s.paragraph_format.space_after=Pt(5)
# header/footer
for section in doc.sections:
    hp=section.header.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.CENTER; rr=hp.add_run(TITLE); set_run_font(rr,'SimSun',8.5,False,(120,120,120))
    fp=section.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); fp._p.append(fld)

# cover
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(90); r=p.add_run('中国产学研合作促进会科技创新奖申报材料'); set_run_font(r,'Microsoft YaHei',20,True,(31,78,121))
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(45); r=p.add_run(TITLE); set_run_font(r,'Microsoft YaHei',26,True,(23,54,93))
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(70)
for line in ['第一完成单位：山西大学','主要完成单位：中国矿业大学；中国电子科技集团第十五研究所；','中国船舶集团第七一九研究所；冀中能源张矿集团；中国舰船研究设计中心','材料类型：申报项目科技创新及产学研合作情况','优化版本：2026年7月（深度扩充版）']:
    r=p.add_run(line+'\n'); set_run_font(r,'Microsoft YaHei',14,False)
doc.add_page_break()

add_heading(doc,'编制说明与材料使用边界',1)
add_p(doc,'本材料在保持原申报书“1、申报项目科技创新—2、申报项目客观评价—3、申报项目产学研合作情况—4、应用情况和效益”总体结构不变的基础上，依据项目技术报告、需求分析、概要设计、应用统计及原申报材料进行系统深化，并结合新型电力系统、配电网高质量发展、人工智能赋能电力运维、无人机自主巡检、多模态融合、边缘智能和预测性维护等公开政策、标准与学术研究进行补充论证。正文继续采用“现实需求—科学问题—关键技术—数学模型—系统实现—客观评价—成果转化—综合效益”的逻辑链条，突出原创性、系统性、工程成熟度、证据可核验性和产学研协同价值。')
add_p(doc,'本版本新增多源时空标定、概率安全约束、感知质量控制、主动学习、长尾损失、跨模态投影、导线悬链线建模、树障动态预测、量测不确定度传播、边云任务卸载、可靠性、贡献分配及经济社会效益等公式；全部公式以Word原生OMML数学对象写入，可在Word中直接选中和编辑。新增15幅系统图均提供独立VSDX源文件，图中方框、文字和连线为可编辑矢量对象；正文逐图引用并说明其逻辑作用。')
add_p(doc,'鉴于现有材料未随附全部第三方检测报告、验收证书、科技查新报告、成果评价证书、专项审计报告、完整知识产权清单和全部用户应用证明，本材料对有关指标继续采用“现有技术事实+证据边界+待补材料”的审慎表述。关键部件识别准确率较主流算法提升5%以上、平台响应时间不高于20 ms、典型缺陷识别精度不低于85%，以及巡检里程、杆塔数量、销售收入和利润等数据，须以最终第三方检测、用户盖章证明和专项审计为正式申报口径。')
add_table(doc,['材料章节','重点优化内容','主要新增成果'],[
['1、科技创新','强化政策背景、学术前沿、科学问题、八项关键技术、公式模型、技术对标和成熟度','公式（1）—（29）；图1—图11、图15；技术对比与评价表'],
['2、客观评价','建立检测、查新、成果评价、可靠性、信息安全和证据映射体系','图12；检测矩阵、查新点、证据链和核验清单'],
['3、产学研合作','完善组织治理、Stage-Gate、接口控制、数据权益、知识产权和收益分配','公式（30）；图13；单位任务与交付物矩阵'],
['4、应用与效益','扩充典型场景、产品化路径、经济/社会/生态效益和推广前景','公式（31）—（35）；图14；效益核算与应用证明模板']],[3.5,8.0,5.0])
doc.add_page_break()

add_heading(doc,'申报项目科技创新及产学研合作情况',1)
add_heading(doc,'1、申报项目科技创新',1)
add_heading(doc,'（1）项目立项背景',2)
add_heading(doc,'1）战略需求与行业发展背景',3)
add_p(doc,'输配电线路是连接电源、主网、配电网和终端用户的关键基础设施，具有点多、线长、面广、跨区域分布、运行环境开放、资产类型复杂和隐患演化速度快等特征。新能源高比例接入、分布式电源快速发展、负荷双向互动以及极端天气频发，使传统基于固定周期和人工经验的巡检模式难以满足新型电力系统对可观、可测、可控、可预测和可追溯的运维要求。国家关于配电网高质量发展、数字电网和人工智能赋能新型工业化的政策导向，持续强调无人巡检终端、状态智能监测、灾害智能感知和设备预测性维护，为本项目提供了明确的战略牵引。')
add_p(doc,'传统人工巡线、登塔检查和人工判图存在安全暴露高、作业周期长、地形依赖强、判据不统一、数据复用不足和隐患发现滞后等问题。无人机虽已显著改善数据采集效率，但大量工程实践仍处于“飞行采集替代人工到场”的阶段，未充分解决GNSS遮挡、弱通信、强风扰动、障碍密集、小目标缺陷、跨季节域偏移、树障空间量测和工单闭环等系统性难题。因此，本项目不是对单一目标检测算法的局部改进，而是面向复杂开放环境构建自主飞行、智能识别、三维量测、风险预测和业务闭环一体化技术体系。')
add_heading(doc,'2）国内外技术进展与尚未解决的问题',3)
add_p(doc,'近年公开研究表明，电力线路自动巡检已由单一CNN检测向多任务视觉、多模态感知、轻量化Transformer、边缘—云协同和自主航空机器人演进。2025年相关综述将研究归纳为部件检测和故障诊断两大主线，并明确提出多模态分析与边云协同是重要发展方向[3]；InsPLAD提供10607幅真实高分辨率无人机图像、17类资产和多类缺陷，揭示了多尺度、多视角、遮挡、透视畸变和复杂背景仍是关键挑战[4]；PLAIR-AW研究进一步说明雾、雨、雪等不利天气会显著降低巡检检测性能[5]。AERIAL-CORE等项目已验证超视距和多类型空中机器人的系统化应用价值，但也指出现有方法往往在受控条件下孤立验证，完整系统集成不足[6]。')
add_p(doc,'在植被管理方面，近期研究已从“识别线路与植被”发展到构造侵限量化指标[7]；2026年的无先验地图杆塔巡检研究利用相机—LiDAR融合实现绝缘子在线定位，说明跨模态语义点云融合是自主巡检的重要方向[8]。然而，公开工作普遍聚焦单一数据集或单一任务，尚缺少将时空标定、鲁棒定位、感知质量约束航迹、少样本识别、树障三维量测、动态风险预测和运维工单贯通的工程化技术链。')
add_table(doc,['技术方向','公开研究进展','仍存不足','本项目响应'],[
['自主巡检','视觉/激光SLAM、感知型MPC、超视距飞行逐步应用','弱GNSS、复杂障碍、成像质量和安全约束难以统一','多源深耦合定位+概率安全约束+感知质量多目标规划'],
['部件与缺陷识别','YOLO、DETR、Transformer和多尺度网络持续提升','小目标、细长结构、遮挡、恶劣天气和跨区域泛化不足','级联注意力+自监督预训练+域泛化+卷积增强Transformer'],
['数据工程','公开数据集推动算法可比性','真实缺陷少、长尾严重、标注昂贵、数据谱系不足','主动学习、困难样本生成、人机复核和版本化数据闭环'],
['树障管理','二维识别、点云提取和侵限指标逐步发展','二维结果难直接给出净空距离，缺乏动态预测和不确定度','图像—LiDAR融合、悬链线拟合、净空置信区间和生长预测'],
['平台应用','边缘推理和云端分析开始融合','模型、告警、工单、证据和持续学习割裂','端—边—云协同、MLOps、工单闭环和申报证据留痕']],[3.0,5.0,5.0,5.0])
add_p(doc,'如图1所示，本项目围绕战略需求和行业痛点，以无人机多源感知为前端，以智能认知和安全决策为核心，以端—边—云平台为工程载体，将应用闭环和成果转化纳入同一创新链。')
add_figure(doc,1,'项目总体技术体系与创新链',fig_paths[1][0],fig_paths[1][1])
add_p(doc,'图2进一步将三类核心科学问题与技术突破及成果形态建立映射，避免创新点之间相互孤立，形成可评价、可查新、可确权和可转化的结构化成果体系。')
add_figure(doc,2,'关键科学问题—技术突破—成果形态映射',fig_paths[2][0],fig_paths[2][1])

add_heading(doc,'（2）具有创新性关键技术内容',2)
add_heading(doc,'1）总体技术路线与关键科学技术问题',3)
add_p(doc,'项目围绕“复杂开放环境下的可靠感知与安全控制”“少样本和域偏移条件下的鲁棒识别”“树障风险的三维量测与动态预测”“多单位成果的工程集成和持续演进”四类问题展开。技术路线采用感知—认知—决策—执行—反馈的闭环架构：多源传感器完成时空一致的数据采集；融合定位与地图构建提供连续状态估计；感知质量约束的航迹规划同时保障安全、能耗和检测可用性；级联注意力、自监督学习和卷积增强Transformer完成关键部件和缺陷识别；图像语义与LiDAR点云融合生成导线—树障数字孪生；端边云平台完成风险告警、工单处置、模型监测和持续训练。')

add_heading(doc,'2）创新点一：多源时空标定与健康度感知的鲁棒融合定位技术',3)
add_p(doc,'针对GNSS遮挡、传感频率不一致、相机与LiDAR外参漂移、IMU零偏累积以及点云局部退化等问题，项目建立时间偏移估计、坐标外参联合标定、误差状态滤波和传感器健康度在线评估方法。通过对GNSS/RTK、IMU、LiDAR、可见光和红外数据进行统一时间戳校正与坐标变换，在单一传感器短时失效时维持连续位姿输出，并向航迹规划模块同步输出状态协方差和可用度。图3给出了多传感器时空标定和鲁棒融合框架。')
add_figure(doc,3,'多传感器时空标定与鲁棒融合定位框架',fig_paths[3][0],fig_paths[3][1])
add_equation(doc,'xₖ = f(xₖ₋₁,uₖ₋₁) + wₖ₋₁',1,'xₖ为时刻k的位姿、速度及传感器偏置状态；uₖ₋₁为IMU或飞控输入；f(·)为非线性状态转移函数；wₖ₋₁为过程噪声。')
add_equation(doc,'zₖˢ = hₛ(xₖ,δtₛ,Tᴮₛ) + vₖˢ',2,'zₖˢ为第s类传感器观测；hₛ(·)为观测模型；δtₛ为时间偏移；Tᴮₛ为传感器坐标系到机体坐标系的外参变换；vₖˢ为观测噪声。')
add_equation(doc,'x̂ₖ|ₖ = x̂ₖ|ₖ₋₁ + Kₖ[zₖ − h(x̂ₖ|ₖ₋₁)]',3,'x̂ₖ|ₖ₋₁和x̂ₖ|ₖ分别为先验与后验状态估计；Kₖ为滤波增益；方括号内为观测创新。')
add_equation(doc,'Pₖ|ₖ = (I − KₖHₖ)Pₖ|ₖ₋₁(I − KₖHₖ)ᵀ + KₖRₖKₖᵀ',4,'P为状态协方差；I为单位矩阵；Hₖ为观测雅可比矩阵；Rₖ为观测噪声协方差。Joseph形式有利于保持协方差数值稳定和半正定性。')
add_equation(doc,'δtₛ* = arg minδt Σᵢ ‖ωᴵᴹᵁ(tᵢ) − ωˢ(tᵢ+δt)‖²',5,'δtₛ*为最优时间偏移；ωᴵᴹᵁ与ωˢ分别为IMU角速度和由第s传感器估计的角速度；i为采样索引。')
add_equation(doc,'Tᴮᴸ* = arg minT Σᵢ ‖π(T pᵢᴸ) − qᵢᴵ‖²Σᵢ⁻¹',6,'Tᴮᴸ为LiDAR到机体/相机基准坐标系的外参；pᵢᴸ为点云特征点；qᵢᴵ为图像对应特征；π(·)为投影函数；Σᵢ为匹配不确定度。')

add_heading(doc,'3）创新点二：感知质量约束的自主巡检与概率安全控制技术',3)
add_p(doc,'项目将传统“最短航程”规划扩展为兼顾安全、能耗、成像视角、运动模糊、目标尺度和任务覆盖度的多目标规划。针对障碍位置和位姿估计均存在不确定性的问题，采用概率安全约束描述碰撞风险，并通过模型预测控制滚动重规划。图4显示任务约束、三维地图、成像质量和概率安全边界共同作用于航迹优化和控制闭环。')
add_figure(doc,4,'感知质量约束的自主飞行与安全控制闭环',fig_paths[4][0],fig_paths[4][1])
add_equation(doc,'P* = arg minP [αL(P)+βR(P)+γE(P)+ηQloss(P)−μC(P)]',7,'P为候选航迹；L为航程；R为碰撞风险；E为能耗；Qloss为成像质量损失；C为任务覆盖度；α、β、γ、η、μ为权重。')
add_equation(doc,'Pr{d(pₖ,Oₖ) ≥ dsafe} ≥ 1−ε,  k=1,…,N',8,'pₖ为预测航迹点；Oₖ为障碍物随机状态；d(·)为空间距离；dsafe为安全距离；ε为允许的最大碰撞概率；N为预测时域。')
add_equation(doc,'Qimg = ω₁Ssharp + ω₂Sillum + ω₃Sscale + ω₄Sangle − ω₅Soccl',9,'Qimg为成像质量综合得分；Ssharp、Sillum、Sscale、Sangle和Soccl分别为清晰度、照度、目标尺度、观测角度和遮挡评分；ω为归一化权重。')
add_equation(doc,'tblur ≤ bpix·Z/(f·vrel)',10,'tblur为允许曝光时间；bpix为允许的像素模糊阈值；Z为目标距离；f为等效焦距；vrel为相机与目标的相对横向速度。该约束用于在近距高速飞行时控制运动模糊。')

add_heading(doc,'4）创新点三：级联注意力—自监督—域泛化的关键部件鲁棒识别技术',3)
add_p(doc,'针对绝缘子、连接金具、线夹、防振锤和导线等目标尺度差异大、局部纹理弱、遮挡与逆光严重的问题，项目将通道注意、空间注意和跨层特征吸收单元级联，使浅层边缘、中层结构和高层语义逐级强化；利用大量未标注巡检影像进行自监督对比预训练，并通过风格扰动、归一化统计随机化和跨域一致性约束提升跨地区、跨季节和跨设备泛化能力。图6展示该网络的双分支学习及融合过程。')
add_figure(doc,6,'级联注意力—自监督—域泛化识别网络',fig_paths[6][0],fig_paths[6][1])
add_equation(doc,'F̃ˡ = φˡ([Fˡ, Ac(Fˡ)⊙Fˡ, As(Fˡ)⊙Fˡ, U(F̃ˡ⁻¹)])',11,'Fˡ为第l层输入特征；Ac和As分别为通道与空间注意；⊙表示逐元素乘法；U为跨层尺度对齐；[·]表示拼接；φˡ为卷积、归一化和非线性映射。')
add_equation(doc,'Lcon = −(1/N)Σᵢ log{exp[sim(qᵢ,kᵢ⁺)/τ]/Σⱼexp[sim(qᵢ,kⱼ)/τ]}',12,'N为样本数；qᵢ为查询特征；kᵢ⁺为同源增强视图正样本；kⱼ为候选键特征；sim为余弦相似度；τ为温度系数。')
add_equation(doc,'Ldg = Σd ‖μd−μ̄‖² + Σd ‖σd−σ̄‖² + λKL(p(y|xᵃ)‖p(y|xᵇ))',13,'d为数据域；μd、σd为域内特征统计；μ̄、σ̄为全域统计；KL为预测分布一致性；xᵃ和xᵇ为不同风格增强视图。')

add_heading(doc,'5）创新点四：面向少样本、长尾与质量漂移的数据工程和主动学习技术',3)
add_p(doc,'项目将数据集建设由一次性标注转化为可持续数据工程。现场数据经过质量筛查、时空配准、自动预标注、人机复核和场景谱系记录，形成可追溯的数据版本；线上低置信度、误检、漏检和新场景样本自动回流，依据不确定度、代表性和多样性进行主动采样；采用Copy-Paste、光照/天气仿真、几何扰动和生成式困难样本扩充长尾类别。图5展示数据、模型与应用的持续训练闭环。')
add_figure(doc,5,'巡检数据治理、主动学习与持续训练闭环',fig_paths[5][0],fig_paths[5][1])
add_equation(doc,'p̃ᵢ = 𝟙[maxc pᵢc ≥ θc]·arg maxc pᵢc',14,'p̃ᵢ为伪标签；pᵢc为样本i属于类别c的预测概率；θc为类别自适应置信阈值；𝟙[·]为指示函数。')
add_equation(doc,'Si = λuH(pᵢ)+λr minj‖zᵢ−zⱼᴸ‖₂+λd D(zᵢ,𝒮)',15,'Si为主动学习采样分数；H为预测熵；zᵢ为样本嵌入；zⱼᴸ为已标注样本嵌入；D为相对当前采样集合𝒮的多样性；λ为权重。')
add_equation(doc,'LFL = −αt(1−pt)ᵞ log(pt)',16,'pt为真实类别预测概率；αt为类别平衡系数；γ为聚焦参数，用于降低易分类样本权重并强化困难样本。')
add_equation(doc,'wc = (1−β)/(1−βⁿᶜ)',17,'wc为类别c的有效样本数权重；nc为该类别样本量；β接近1，用于缓解长尾类别不平衡。')
add_equation(doc,'Ltotal = λclsLcls+λboxLbox+λobjLobj+λconLcon+λdgLdg+λreg‖θ‖²₂',18,'Lcls、Lbox、Lobj分别为分类、定位和目标置信度损失；Lcon和Ldg为自监督与域泛化损失；θ为模型参数；各λ为损失权重。')

add_heading(doc,'6）创新点五：卷积增强Transformer的多尺度缺陷检测与结构解析技术',3)
add_p(doc,'标准CNN在局部纹理提取方面具有优势，但长距离结构关系建模不足；纯Transformer具有全局建模能力，却可能弱化裂纹、锈蚀、断股和小型缺件等局部细节。项目将深度可分离卷积构成的局部窗口连接模块（LWC）和跨窗口卷积连接模块（GWC）嵌入Patch Embedding、Patch Merging和窗口Transformer块，并引入跨层特征金字塔、轻量检测头和分割/关键点辅助任务，实现局部细节、全局语义和几何结构的协同建模。图7给出详细架构。')
add_figure(doc,7,'卷积增强Transformer缺陷检测架构',fig_paths[7][0],fig_paths[7][1])
add_equation(doc,'Attention(Q,K,V)=softmax(QKᵀ/√dk+B)V',19,'Q、K、V分别为查询、键和值矩阵；dk为键向量维度；B为相对位置偏置或窗口掩码。')
add_equation(doc,'Yˡ = Xˡ + MSA[LN(Xˡ)] + CLG[LN(Xˡ)]',20,'Xˡ为第l个块输入；MSA为窗口或移位窗口多头注意力；LN为层归一化；CLG为LWC/GWC卷积增强模块；Yˡ为中间特征。')
add_equation(doc,'Xˡ⁺¹ = Yˡ + MLP[LN(Yˡ)]',21,'MLP为前馈网络；残差连接维持梯度传播并促进全局语义与局部结构融合。')
add_equation(doc,'Foutˡ = Σj wij·Resize(Fj),  wij=exp(aij)/Σmexp(aim)',22,'Fj为不同尺度输入特征；Resize为尺度对齐；wij为可学习归一化融合权重；aij为未归一化权重。')

add_heading(doc,'7）创新点六：图像语义—LiDAR点云融合的导线与树障空间量测技术',3)
add_p(doc,'项目首先利用分割网络提取导线、树冠和杆塔语义，再依据相机—LiDAR外参和时间同步结果将语义投影到三维点云；结合IMU位姿补偿、点云去噪、导线点集拟合和树冠聚类，得到可解释的档距级三维对象。图8所示流程将二维语义与三维几何相结合，避免仅依赖单目尺度估计造成的距离歧义。')
add_figure(doc,8,'图像语义—LiDAR点云跨模态融合流程',fig_paths[8][0],fig_paths[8][1])
add_equation(doc,'ũ = K·Tᶜᴸ·[pᴸ;1],  (u,v)=(ũ₁/ũ₃,ũ₂/ũ₃)',23,'pᴸ为LiDAR点；Tᶜᴸ为LiDAR到相机的外参；K为相机内参；ũ为齐次像素坐标；(u,v)为图像坐标。')
add_equation(doc,'z(x)=a cosh((x−x₀)/a)+c',24,'z(x)为导线在局部垂直平面中的高度；a为悬链线尺度参数；x₀为最低点水平位置；c为高度偏置。该模型用于抑制离散点云噪声并估计弧垂。')
add_equation(doc,'dmin = minp∈PL,q∈PT ‖p−q‖₂',25,'PL为导线拟合点集；PT为树冠点集；p和q分别为导线点与树冠点；dmin为实测最小净空距离。')
add_equation(doc,'dsafe = dmin−Δpose−Δfit−Δwind−Δgrowth',26,'Δpose为位姿与配准误差裕度；Δfit为导线拟合误差；Δwind为风偏和弧垂裕度；Δgrowth为预测周期内的树木生长裕度；dsafe为保守净空距离。')
add_equation(doc,'σd² ≈ JpΣpJpᵀ + JqΣqJqᵀ + σcal²',27,'Σp和Σq分别为导线点和树冠点协方差；Jp、Jq为距离函数雅可比；σcal²为标定误差方差；σd为净空量测标准不确定度。')

add_heading(doc,'8）创新点七：树障数字孪生、动态预测与可解释风险分级技术',3)
add_p(doc,'树障风险并非静态距离问题，而是受树种生长、季节、降雨、风偏、导线温度弧垂和倒伏概率共同影响。项目以档距为基本单元构建树障数字孪生，将历史点云、气象、植被结构和处置记录关联，预测未来净空距离分布，并生成提示、一般、严重和紧急四级风险。图9显示从空间基准、状态估计、未来概率分布到分级处置的全过程。')
add_figure(doc,9,'树障数字孪生、动态预测与分级处置',fig_paths[9][0],fig_paths[9][1])
add_equation(doc,'hₜ₊Δ = hₜ + r(s,m)Δ + 0.5aₕΔ² + εₜ',28,'hₜ和hₜ₊Δ为当前及预测时刻树冠高度；r(s,m)为由树种s和月份m决定的生长率；aₕ为生长加速度项；εₜ为过程扰动。')
add_equation(doc,'Rveg = σ[a₁(Dv−E[d])/Dv + a₂Pr(d<Dv)+a₃g+a₄c+a₅u]',29,'Rveg为树障风险指数；Dv为电压等级对应控制阈值；E[d]为预测净空均值；Pr(d<Dv)为越限概率；g为生长速率；c为倒伏/倾斜结构风险；u为量测不确定度；a为权重；σ为Sigmoid函数。')

add_heading(doc,'9）创新点八：端—边—云协同在线平台与可信MLOps技术',3)
add_p(doc,'项目平台支持可见光、红外、视频、LiDAR点云、GNSS/INS和气象参数等多源数据接入。边缘侧完成质量筛查、压缩加密、快速推理和异常优先上报；云端完成高精度推理、GIS时空分析、资产关联和跨线路模型训练；运维侧完成告警确认、任务派发、工单处置和结果回填。通过模型注册、数据版本、灰度发布、性能漂移监测、回滚和审计日志建立可信MLOps闭环。图10展示任务卸载和模型持续演进关系。')
add_figure(doc,10,'端—边—云协同推理与MLOps闭环',fig_paths[10][0],fig_paths[10][1])
add_equation(doc,'xᵢ* = arg minxᵢ∈{0,1} [xᵢ(Tedgeᵢ+Eedgeᵢ)+(1−xᵢ)(Ttxᵢ+Tcloudᵢ+λEtxᵢ)]',30,'xᵢ=1表示任务i在边缘执行，xᵢ=0表示卸载至云端；Tedge、Ttx、Tcloud分别为边缘推理、传输和云端推理时延；Eedge、Etx为能耗；λ为能耗权重。')
add_equation(doc,'Tend = Tacq+Tpre+Tinf+Tpost+Ttx+Tbiz',31,'Tend为端到端响应时间；Tacq、Tpre、Tinf、Tpost、Ttx、Tbiz分别为采集、预处理、推理、后处理、传输和业务入库时延。20 ms指标必须明确是否包含上述各阶段。')
add_equation(doc,'A = MTBF/(MTBF+MTTR)',32,'A为系统可用度；MTBF为平均无故障时间；MTTR为平均修复时间。该指标用于评价平台连续运行和工程可靠性。')

add_heading(doc,'10）创新点的系统集成关系与技术成果形态',3)
add_p(doc,'八项创新相互依赖：鲁棒定位和感知质量控制为高质量影像与点云提供前提；数据闭环、自监督和域泛化共同提升识别模型跨场景能力；卷积增强Transformer输出缺陷、导线和树冠语义，为三维量测提供先验；树障数字孪生将量测结果转化为可解释风险；端边云平台承载模型、告警和工单，现场结果再次回流数据闭环。图11从单一人工巡检、单任务检测和离线多模态分析出发，给出本项目向系统样机、示范应用和标准化产品演进的技术成熟度路径。')
add_figure(doc,11,'技术先进性对标与成熟度提升路径',fig_paths[11][0],fig_paths[11][1])
add_table(doc,['创新点','核心突破','主要指标/评价方法','成果形态'],[
['多源时空标定与鲁棒定位','时间偏移、外参、误差状态和健康度联合估计','定位误差、姿态误差、失锁恢复时间、协方差一致性','融合导航算法、标定工具、嵌入式模块'],
['感知质量自主飞行','安全、能耗、成像质量和覆盖度联合优化','航迹偏差、避障成功率、任务完成率、有效影像率','规划控制软件、飞行试验报告'],
['级联注意力与域泛化','局部—跨层—全局特征和未标注数据联合利用','mAP、Recall、小目标AP、跨域性能下降率','识别模型、论文、专利'],
['数据工程与主动学习','低置信度回流、长尾平衡、版本和谱系管理','标注节约率、增量样本收益、数据完整性','专用数据集、标注工具、数据规范'],
['卷积增强Transformer','局部卷积与窗口全局建模协同','缺陷mAP/F1、参数量、FLOPs、端侧时延','检测模型、SDK、软著'],
['图像—LiDAR量测','语义投影、悬链线拟合和净空不确定度','MAE、RMSE、95%置信区间覆盖率','量测算法、标定报告'],
['数字孪生风险预测','动态净空、越限概率和处置优先级','预警提前量、风险AUC、处置闭环率','树障风险模块、企业规范'],
['端边云与MLOps','任务卸载、模型注册、漂移监测和回滚','端到端时延、可用度、并发、模型更新周期','在线平台、接口标准、应用示范']],[3.2,5.0,5.3,4.4])

add_heading(doc,'（3）主要应用范围',2)
add_p(doc,'项目成果适用于500 kV、220 kV、110 kV及10 kV等不同电压等级的输配电线路，可根据现行标准和企业规程配置净空阈值、巡检任务、风险等级和复核周期。应用对象包括杆塔、导线、地线、绝缘子串、连接金具、线夹、防振锤、间隔棒等关键部件，以及异物附挂、缺件、破损、锈蚀、松动、发热、断股、树木侵限和倒伏风险等缺陷隐患。')
add_p(doc,'在空间场景方面，适用于山区陡坡、林区密植、城镇建筑密集区、矿区线路走廊、跨河跨谷区、弱通信区及逆光、雾霾、雨雪和季节变化明显的复杂环境；在业务场景方面，适用于周期巡检、精细化复检、灾后特巡、树障专项治理、工程验收、通道数字化建模和状态评价；在产业场景方面，可向电网企业、发电企业自备线路、矿山和园区供电系统、无人机装备制造商、边缘计算厂商和智能运检服务商推广。')
add_table(doc,['应用场景','核心功能','典型部署','应用价值'],[
['高压输电线路','长距离航巡、杆塔精细巡检、树障和异物排查','固定翼/多旋翼+机载LiDAR+云平台','提高覆盖效率，降低登塔和偏远地区作业风险'],
['城镇配电网','部件缺陷、热异常、异物和树障在线识别','多旋翼+固定终端+边缘节点','缩短发现与处置时间，支撑精益运维'],
['山区/林区通道','弱GNSS定位、导线跟随、避障和净空量测','LiDAR+INS/GNSS+嵌入式计算','解决交通不便、遮挡和弱通信作业难题'],
['矿区与工业园区','高粉尘复杂地形和关键供电走廊巡检','无人机+专网+企业运维接口','保障连续生产，减少非计划停电'],
['灾害应急','冰雪、强风、山火、洪涝后快速复查','快速部署无人机/多机协同+应急平台','快速形成灾情态势和抢修优先级'],
['装备与服务产业','算法SDK、边缘终端、平台软件和巡检服务','模块化产品、许可授权和联合解决方案','形成可复制产品和规模化服务收入']],[3.0,5.2,4.6,5.0])

add_heading(doc,'（4）技术水平',2)
add_p(doc,'项目技术水平从原创性、系统性、先进性、工程成熟度、可信性和可推广性六个维度评价。原创性体现在时空标定、感知质量航迹、级联注意力、自监督与域泛化、卷积增强Transformer、语义点云量测、动态风险预测和可信MLOps的组合创新；系统性体现在从感知、控制、识别、量测到工单的全链条；先进性体现在解决小目标、恶劣天气、跨域泛化、多模态融合和边缘部署前沿难题；工程成熟度体现在在线平台和典型场景试点；可信性体现在不确定度、模型健康度、审计日志和证据链；可推广性体现在模块化接口和阈值配置。')
add_table(doc,['评价维度','项目水平表述','建议量化指标','必须补充的证据'],[
['识别性能','关键部件准确率较同条件主流算法提升5%以上；典型缺陷识别精度不低于85%','Precision、Recall、F1、mAP50-95、Small-AP、跨域下降率','固定测试集、基线配置、随机种子、模型版本和第三方复测'],
['自主飞行','具备三维建图、深耦合定位、导线跟随、质量约束规划和避障能力','位置/姿态RMSE、航迹偏差、避障成功率、有效影像率、任务完成率','架次记录、飞控日志、视频和安全评估'],
['树障量测','形成档距级三维净空、置信区间和动态分级能力','MAE、RMSE、95%区间覆盖、预警提前量','全站仪/基准点云比对、标定和现场复测'],
['平台性能','形成端边云协同和闭环处置；目标响应时间不高于20 ms','端到端分项时延、并发、可用度、连续运行时间','硬件、输入尺寸、批次、网络条件和原始日志'],
['工程应用','覆盖多电压等级、多地形和多季节场景','里程、杆塔、缺陷发现率、闭环率、复检一致率','用户盖章证明、任务日志、验收材料'],
['成果转化','形成算法、终端、平台和服务组合','合同收入、利润、许可、标准、培训和市场覆盖','合同、发票、财务账和专项审计']],[3.0,5.5,5.0,4.5])
add_p(doc,'综合现有材料，项目具有申报“整体达到国内领先水平，部分关键技术达到国际先进水平”的技术基础，但最终表述必须以具有资质的第三方检测、科技查新和科技成果评价结论为依据。申报书中应避免未经证据支撑的“国际领先”“国内首创”等绝对性描述。')

add_heading(doc,'（5）主要开发历程',2)
add_p(doc,'项目遵循“需求牵引—科学问题凝练—核心算法攻关—系统集成—现场验证—第三方评价—产品定型—规模推广”的阶段路线。图15将原有开发过程扩展为Stage-Gate阶段门管理，并强调现场问题持续回流数据和模型。')
add_figure(doc,15,'项目开发历程与持续演进路线',fig_paths[15][0],fig_paths[15][1])
add_table(doc,['阶段','主要任务','阶段成果','阶段门/证据'],[
['需求分析与方案论证','用户访谈、场景分类、风险分析、指标定义和总体架构','需求分析、概要/详细设计、技术路线','G0：立项批复、需求确认、专家论证和合作协议'],
['核心算法攻关','多源融合定位、级联注意力、自监督、域泛化、主动学习','算法原型、数据集、论文专利和对比试验','G1：代码版本、训练日志、数据合规和实验复现'],
['缺陷与树障融合感知','卷积增强Transformer、点云融合、悬链线和风险模型','检测模型、量测算法、风险分级原型','G2：标定报告、测距比对、消融和鲁棒性测试'],
['平台集成与边缘部署','前端、边缘、通信、GIS、MLOps和安全机制集成','在线平台、边缘模块、接口和软著','G3：系统测试、网络安全、软件测评和用户手册'],
['现场试点与第三方评价','多电压等级、多地形和多季节应用验证','试点应用、问题清单、检测和成果评价','G4：飞行记录、用户证明、CMA/CNAS报告和查新'],
['产品定型与规模推广','产品认证、标准化、市场和区域交付能力建设','SDK、终端、平台、服务和销售收入','G5：合同、发票、专项审计、标准和推广协议']],[2.8,5.6,4.8,5.0])

add_heading(doc,'2、申报项目客观评价',1)
add_heading(doc,'（1）技术检测报告、验收报告、鉴定报告',2)
add_p(doc,'客观评价必须覆盖算法、系统、飞行、量测、可靠性、网络安全和应用效果，形成“原始数据—内部测试—第三方检测—验收/评价—申报事实”的逐级证据链。图12所示证据链既服务奖项申报，也服务后续产品认证和成果转化。')
add_figure(doc,12,'客观评价与申报证据链',fig_paths[12][0],fig_paths[12][1])
add_table(doc,['报告类型','核心检测内容','统计要求','规范引用表述（待补证）'],[
['算法性能检测','关键部件、典型缺陷和树障识别；小目标与复杂天气分层','置信区间、重复试验、混淆矩阵、跨域测试','“经××机构检测，模型在××测试集的××指标为××。”'],
['自主飞行试验','定位、跟随、避障、有效影像、任务完成和失联保护','不同地形/天气/通信条件、架次和失败原因','“在×类场景开展×架次试验，任务完成率为××%。”'],
['树障量测检测','标定、配准、悬链线、净空和风险等级','按距离、姿态、树冠密度分层，报告MAE/RMSE/区间覆盖','“与全站仪/基准点云相比，净空MAE为××m。”'],
['软件与平台测评','功能、并发、时延、稳定、容错、接口、审计和模型管理','明确硬件、输入、网络、连续运行时间','“平台连续运行××h，无严重故障，可用度为××。”'],
['网络与数据安全','身份认证、最小权限、加密、日志、脱敏和恢复','渗透测试、漏洞扫描、备份恢复演练','“经安全测评，满足××等级和企业安全要求。”'],
['验收/成果评价','任务完成、创新性、成熟度、应用和推广价值','专家组成、回避机制、证据完整性','“专家组认为项目完成任务书内容，整体达到××水平。”']],[3.0,5.3,5.0,5.0])
add_heading(doc,'（2）查新报告',2)
add_p(doc,'建议委托具有资质的科技查新机构，检索Web of Science、Scopus、IEEE Xplore、Engineering Village、中国知网、万方、国家知识产权局和WIPO。查新点应描述具体结构、输入输出、约束和区别特征，不宜仅写“采用人工智能”“采用无人机”等宽泛组合。')
add_table(doc,['建议查新点','区别于常规方案的核心特征','建议检索要素'],[
['多源时空标定与感知质量自主飞行','时间偏移、外参、协方差健康度、概率安全和成像质量联合约束','UAV power line, temporal calibration, chance constraint, perception-aware planning'],
['级联注意力—自监督—域泛化识别','通道/空间/跨层注意与未标注预训练、风格扰动和主动学习联合','cascade attention, self-supervised, domain generalization, power line asset'],
['卷积增强Transformer缺陷检测','Patch阶段嵌入LWC/GWC并融合多尺度检测与几何辅助任务','convolution-enhanced transformer, shifted window, defect detection'],
['图像—LiDAR树障空间量测','语义投影、悬链线拟合、净空不确定度、风偏与生长裕度','camera LiDAR fusion, catenary, vegetation clearance, uncertainty'],
['树障数字孪生和动态风险闭环','档距级时序孪生、越限概率、分级告警、工单与模型回流','digital twin, vegetation growth, probabilistic risk, work order'],
['端边云可信MLOps','任务卸载、模型注册、漂移监测、灰度发布、回滚和审计','edge cloud collaboration, MLOps, model drift, power inspection']],[4.2,7.2,6.8])
add_heading(doc,'（3）国内外同行公开学术评价与技术趋势',2)
add_p(doc,'同行研究总体趋势与本项目技术路线具有高度一致性：公开综述强调边云协同和多模态分析[3]；InsPLAD和TTPLA表明高分辨率图像中的尺度变化、细长导线、遮挡、复杂背景和缺陷稀缺仍有较大提升空间[4,9]；PLAIR-AW说明不利天气恢复与检测需要联合考虑[5]；AERIAL-CORE显示自主航空机器人可显著降低人员风险并提高长距离巡检能力[6]；植被管理研究开始构建侵限量化指标[7]；相机—LiDAR融合的无先验地图巡检进一步验证了在线语义定位价值[8]。本项目在此基础上实现“自主采集—鲁棒识别—三维量测—动态预测—工单处置—模型迭代”的系统集成，具有明显的工程完整性。')
add_table(doc,['公开研究方向','代表性进展','本项目对应提升'],[
['多模态与边云协同','成为自动电力线路巡检的重要发展方向','将多模态识别、空间量测、风险和业务闭环统一实现'],
['高分辨率资产数据集','推动多尺度资产、缺陷和异常检测','建设国内复杂场景数据闭环并记录场景谱系和质量漂移'],
['感知型自主飞行','将路径控制与观测质量联合优化','叠加概率安全、深耦合定位、导线跟随和有效影像指标'],
['轻量YOLO/DETR/Transformer','持续强化小目标、边缘和复杂背景','级联注意力+自监督+域泛化+CET多尺度检测'],
['植被侵限检测','从目标识别向侵限指标演进','LiDAR直接量测净空，加入不确定度、悬链线、风偏和生长预测'],
['在线工程系统','长距离自主系统开始验证','增加MLOps、工单、证据链、产品化和产学研治理']],[4.0,6.5,7.5])
add_heading(doc,'（4）其他评价与成果支撑',2)
add_p(doc,'根据现有项目简介，项目已形成30余项中国发明专利、6项软件著作权和80余篇论文等成果线索，并在2023—2025年形成销售和利润数据。正式申报应按照“与本项目直接相关、署名单位一致、贡献关系明确、无权属争议、未重复使用”的原则遴选，建立成果—创新点—完成单位—完成人—应用证据五维对应表。')
add_table(doc,['成果类别','申报呈现方式','核验重点'],[
['发明专利','名称、专利号、授权日、权利人、创新点和实施情况','授权状态、技术关联、权属和重复报奖'],
['软件著作权','软件名称、登记号、版本、功能和应用单位','是否对应平台、边缘算法、数据或量测工具'],
['论文/专著','优先列出高质量、高被引且直接支撑核心技术的代表作','检索收录、作者单位、项目标注和他引评价'],
['标准/规范','已发布、在编的团体、企业或行业标准及纳入条款','标准状态、排名和技术贡献'],
['应用证明','用户、线路、电压、里程、杆塔、时间、效果和联系人','盖章、日志可追溯、与合同和审计一致'],
['经济效益','销售、服务、成本节约、利润和避免损失','合同、发票、财务账、专项审计和剔除非项目收入']],[3.2,7.5,7.0])

add_heading(doc,'3、申报项目产学研合作情况',1)
add_heading(doc,'（1）申报项目的产学研合作情况',2)
add_p(doc,'项目由山西大学作为第一完成单位，联合中国矿业大学、中国电子科技集团第十五研究所、中国船舶集团第七一九研究所、冀中能源张矿集团和中国舰船研究设计中心，形成高校、科研院所和行业用户协同的“产学研用”创新联合体。合作并非简单成果拼接，而是以同一总体架构、同一数据基线、同一接口规范、同一评价体系和同一转化目标组织任务。图13展示Stage-Gate阶段门、技术委员会、应用示范组和知识产权转化组之间的治理关系。')
add_figure(doc,13,'产学研Stage-Gate协同治理与知识产权机制',fig_paths[13][0],fig_paths[13][1])
add_table(doc,['完成单位','建议主要职责与贡献定位','主要交付物和贡献证据'],[
['山西大学（第一完成单位）','总体设计与组织协调；级联注意力、自监督/域泛化、数据闭环、卷积增强Transformer、风险模型和成果凝练','总体方案、算法、数据规范、论文专利、平台核心模块和管理材料'],
['中国矿业大学','复杂地形和矿区场景建模；点云处理、自主导航和路径规划；矿区线路验证','点云/导航算法、仿真与试验报告、场景数据和应用证明'],
['中国电子科技集团第十五研究所','边缘计算、软件工程、通信与数据安全、平台架构和接口集成','边缘软件、平台模块、接口规范、软件和安全测评'],
['中国船舶集团第七一九研究所','多传感器嵌入式集成、复杂环境可靠性、半实物仿真和工程验证','嵌入式原型、可靠性试验、硬件接口和验证报告'],
['冀中能源张矿集团','提出矿区运维需求，提供场景与业务流程，开展现场试验、示范和效益评价','需求确认、现场数据、试点线路、用户证明和效益材料'],
['中国舰船研究设计中心','系统总体集成、任务规划、复杂系统工程评估和产品化设计','集成方案、任务规划模块、工程评价和产品化建议']],[3.5,8.2,6.0])
add_p(doc,'上述分工应以正式合作合同、任务书、知识产权协议、论文专利署名、代码提交、试验记录和用户证明为最终依据，不得以单位专业优势替代真实贡献证明。')
add_heading(doc,'合作目标、模式和运行机制',3)
for b in [
'合作目标：形成具有自主知识产权、可工程部署、可复制推广的输配电线路无人机巡检和树障智能感知成套技术，包括自主飞行算法、专用数据集、识别模型、树障量测与风险模块、端边云平台、标准规范和示范应用。',
'合作模式：实行“用户出题—联合论证—高校与科研院所协同攻关—用户现场验证—第三方评价—共同转化”，用户需求和真实场景贯穿研发全过程。',
'技术治理：设置项目决策委员会和技术委员会，执行需求冻结、算法原型、系统联调、现场测试、第三方评价和推广应用阶段门；跨单位接口实行“接口责任人+基线文档+联调清单”。',
'数据治理：原始影像、点云、台账和日志分级分类管理，明确采集授权、脱敏、使用、留存、共享和销毁流程；联合训练数据记录来源、许可和版本谱系。',
'质量机制：算法同时接受标准数据/实验室对比和真实线路测试，既评价精度，也评价时延、资源、可靠性、可维护性和安全性；失败样本全部进入问题库。',
'人才机制：研究生、青年教师和工程技术人员交叉驻场、联合指导和共同测试，形成理论方法、工程经验和行业知识双向流动。']:
    add_bullet(doc,b)
add_heading(doc,'知识产权与利益分配',3)
add_p(doc,'知识产权和成果收益坚持“合同约定优先、实际贡献为基础、背景知识产权与项目新增成果分离、技术原创与持续服务兼顾”的原则。项目启动时列明各方背景知识产权；新增专利、软件著作权、数据集和技术秘密按发明贡献、开发投入、试验验证、项目管理和市场转化确定权属和署名；对外许可、产品销售和技术服务收益，在扣除必要成本后按协议分配。建议采用量化贡献系数作为谈判辅助工具：')
add_equation(doc,'πᵢ = (Σr ωr·cᵢr)/(ΣjΣr ωr·cjr)',33,'πᵢ为第i个单位的建议收益分配比例；cᵢr为单位i在第r类贡献中的确认分值；ωr为技术原创、软硬件开发、数据试验、管理和转化等贡献权重；最终比例以各方签署协议为准。')

add_heading(doc,'（2）管理创新、机制创新方面的经验与做法',2)
add_p(doc,'项目形成了以场景定义科研问题、以数据和接口组织协作、以双验证保证质量、以失败样本驱动持续创新、以知识产权地图支撑转化、以贡献留痕支撑奖项和收益分配的管理模式。其核心是把多单位专业优势转化为同一系统的可交付能力，而不是在项目末期进行模块拼接。')
add_table(doc,['机制创新','具体做法','形成的可推广经验'],[
['场景牵引','用户先描述任务、风险、处置流程和评价指标，技术团队再凝练科学问题','避免算法与业务脱节，指标可直接进入检测和验收'],
['接口基线','统一时间、坐标、数据格式、置信度、模型版本和日志接口','降低跨单位集成成本，支持模块替换和复用'],
['双验证','实验室标准对比+真实线路现场测试','兼顾论文指标与工程可靠性'],
['失败样本闭环','误检、漏检、模糊、点云异常和定位漂移进入问题库','形成问题—数据—算法—验证的持续创新循环'],
['知识产权地图','围绕定位、识别、训练、量测和平台形成专利组合','明确背景技术、共有成果和可许可产品模块'],
['贡献留痕','任务书、代码、试验、会议、论文专利和应用证明关联存档','支撑完成单位/完成人排序、收益分配和奖项核查']],[3.0,7.5,7.0])

add_heading(doc,'（3）科技成果转化及产学研合作对接需求',2)
add_p(doc,'项目成果适合形成“1个平台+3类产品+N类场景服务”：1个平台为输配电线路无人机巡检与树障隐患智能感知平台；3类产品包括机载/地面边缘智能终端、算法SDK与模型包、树障量测与风险管理模块；N类服务包括线路建模、周期巡检、灾后特巡、树障专项治理、缺陷复核、数据治理和模型定制。')
add_table(doc,['对接方向','具体需求','潜在对象','合作方式'],[
['工程化与认证','硬件定型、环境适应、电磁兼容、可靠性、网络安全和软件测评','无人机/载荷厂商、检测认证机构、工业计算企业','联合开发、委托检测、产品认证'],
['规模场景验证','跨区域、季节、电压等级和地形的长期运行','电网、发电、矿山、园区和大型能源集团','示范项目、联合实验室、场景开放'],
['标准与规范','数据格式、标注、算法评价、净空量测、告警和接口规范','行业协会、标准化机构、电网科研院所','团体/企业标准和行业标准预研'],
['数据和模型生态','合规数据共享、联邦训练和模型评测','高校、科研院所、行业用户和数据服务商','数据合作、联合评测、隐私计算'],
['产业化与市场','区域交付、运维、培训和售后能力','集成商、无人机服务商和产业基金','许可、作价入股、联合投标、转化基金'],
['国际合作','海外山地、林区和长距离线路适配','国际电力企业、承包商和研究机构','联合研发、国际标准和海外示范']],[3.0,5.2,5.0,4.5])

add_heading(doc,'4、应用情况和效益',1)
add_heading(doc,'（1）推广应用及典型案例',2)
add_p(doc,'根据项目简介，成果已在电网公司、线路运维单位、矿山企业和智能巡检装备企业开展试点，覆盖110 kV、220 kV和500 kV等电压等级；材料汇总口径为累计巡检线路超过2000 km、覆盖杆塔超过5000基、巡检成本较传统模式降低50%以上、自主巡检效率达到人工巡检10倍以上、单基杆塔巡检时间由约1 h缩短至5 min以内。上述数据均须对应用户盖章应用证明、任务日志和统计口径。')
add_p(doc,'图14展示成果由算法和数据资产向边缘产品、在线平台、巡检服务和多行业场景扩散，并最终形成效率、风险、产业和综合效益的传导链。')
add_figure(doc,14,'成果产品化、场景复制与综合效益传导',fig_paths[14][0],fig_paths[14][1])
add_table(doc,['典型案例/场景','应用内容','可量化成效','应附证明'],[
['冀中能源张矿集团矿区线路（重点案例）','复杂地形自主飞行、部件/树障识别、风险告警和工单闭环','巡检里程、杆塔数、缺陷数、闭环率、时间和成本变化','盖章应用证明、任务日志、现场照片、验收和效益说明'],
['山区/林区输电线路','导线跟随、LiDAR避障、净空量测和生长预测','有效影像率、量测误差、预警提前量和树障处置率','飞行记录、全站仪比对、树障工单'],
['城镇配电网','小目标部件缺陷、异物和热异常识别','识别精度、响应时延、发现到处置时间','平台日志、缺陷台账和复检结果'],
['灾害应急特巡','强风、冰雪、山火或洪涝后快速态势获取','覆盖速度、异常定位时间和抢修优先级准确性','应急任务单、飞行轨迹和抢修记录'],
['智能巡检产品配套','SDK、边缘终端和平台授权','产品数量、客户数量、收入和续费/服务','合同、发票、产品检测和用户证明']],[3.5,6.0,5.0,5.0])

add_heading(doc,'（2）直接经济效益',2)
add_p(doc,'现有项目简介给出的2023—2025年累计新增销售额为97607.8万元、累计新增利润为32012.5万元。正式申报前必须由专项审计明确统计期间、收入确认原则、项目产品边界、关联交易、成本分摊及非项目收入剔除。经济效益建议同时从新增收入、成本节约、避免停电损失和净现值四个维度核算。')
add_equation(doc,'Bdirect = ΔS + Csave + Lavoid − Cinc',34,'Bdirect为直接经济效益；ΔS为项目新增销售或技术服务收入；Csave为巡检、人工、车辆和外委成本节约；Lavoid为避免故障停电和设备损坏损失；Cinc为因部署新增的设备、通信、运维和折旧成本。')
add_equation(doc,'NPV = Σt=0ᵀ (Bt−Ct)/(1+r)ᵗ',35,'NPV为项目净现值；Bt和Ct分别为第t年的收益和成本；r为折现率；T为评价期。')
add_equation(doc,'Lavoid = Σe pe·[Pload,e·tout,e·Ve + Crepair,e + Csecondary,e]',36,'e为风险事件类型；pe为采用项目后被提前发现并避免的事件概率；Pload,e为影响负荷；tout,e为停电时间；Ve为单位电量损失价值；Crepair和Csecondary为抢修与次生损失。')
add_table(doc,['效益项目','计算口径','原始凭证','审计注意事项'],[
['产品/软件销售','与项目直接相关的终端、平台、SDK和模型包收入','合同、发票、收款、出库/验收','剔除其他产品和关联交易非公允部分'],
['技术与巡检服务','项目实施、定制、巡检和运维服务收入','服务合同、任务单、验收和工时','按项目实际完成比例确认'],
['成本节约','人工、车辆、登塔、外委、差旅和重复巡检减少','成本台账、人员/车辆记录和历史对比','统一基准年和工作量口径'],
['避免损失','提前识别缺陷/树障避免的停电和设备损失','缺陷工单、故障概率、负荷和修复成本','采用审慎概率，不把潜在最大损失全部计入'],
['新增利润','项目收入扣除直接和合理分摊成本','财务账、成本分摊和所得税口径','专项审计说明计算方法和边界']],[3.2,6.0,5.0,5.0])

add_heading(doc,'（3）社会、安全与生态效益',2)
add_p(doc,'项目通过无人机替代登塔、陡坡徒步和近电作业，减少人员在高处、带电、恶劣天气和复杂地形中的暴露；通过高频巡检和风险预测提前发现缺陷与树障，提升供电可靠性和应急处置能力；通过减少车辆、人工往返和重复飞行降低能源消耗与碳排放；通过数据和模型沉淀促进电力运维数字化、复合型人才培养和相关装备产业发展。')
add_equation(doc,'ΔC = Σm (Qbase,m−Qproj,m)·EFm',37,'ΔC为减排量；m为车辆燃油、电力消耗或其他能源类型；Qbase,m和Qproj,m分别为基准巡检和项目巡检的能源消耗；EFm为对应排放因子。该指标应依据可核验里程、油耗/电耗和权威排放因子计算。')
add_equation(doc,'Rsafety = 1 − (Hproj·Eproj)/(Hbase·Ebase)',38,'Rsafety为安全暴露降低率；H为人员作业工时；E为单位工时风险暴露系数；base和proj分别代表传统模式和项目模式。风险系数应由企业安全管理部门依据作业类型核定。')
add_table(doc,['效益维度','主要表现','建议指标'],[
['安全效益','减少登塔、近电、陡坡和恶劣天气作业','高风险工时减少率、无人化替代率、事故/未遂事件变化'],
['供电可靠性','提前发现缺陷和树障，缩短发现—处置周期','隐患发现率、闭环率、平均处置时间、避免停电事件'],
['应急能力','灾后快速获取线路态势和抢修优先级','单位时间覆盖里程、异常定位时间、复电时间'],
['生态效益','减少车辆和重复巡检，精细化树障修剪','燃油/电耗、碳减排、精准修剪率和植被影响面积'],
['人才与产业','培养AI、电力、无人机和系统工程复合人才','联合培养人数、培训人次、标准、专利和产业合作项目']],[3.0,8.0,7.0])

add_heading(doc,'（4）推广前景与持续发展计划',2)
add_p(doc,'项目具有从单线路试点向区域平台、从单一电压等级向多电压等级、从电网巡检向矿山和园区自备线路、从一次性交付向持续模型服务扩展的条件。后续重点包括：完成第三方检测和产品认证；在不同地区、季节和设备上进行长期稳定性验证；形成数据、标注、算法评价、树障量测和平台接口标准；建立区域交付和运维体系；探索联邦学习和隐私计算以扩大数据生态；围绕重点行业形成模块化联合解决方案。')
add_table(doc,['阶段','推广目标','关键行动','风险控制'],[
['近期（1年）','完成证据补强和产品定型','第三方检测、查新、评价、网络安全和可靠性测试','指标口径统一、知识产权和数据合规'],
['中期（2—3年）','跨区域、多场景规模应用','建设示范线、区域交付中心、标准和培训体系','模型漂移监测、备件和服务能力'],
['远期（3—5年）','形成行业生态和国际合作','平台开放接口、联合解决方案、标准和海外适配','国际法规、数据跨境和本地化运维']],[3.0,5.0,6.5,4.5])

add_heading(doc,'附录A 主要公式与变量索引',1)
formula_rows=[]
formula_desc={1:'无人机非线性状态转移',2:'多传感器观测与时空参数',3:'滤波状态更新',4:'协方差稳定更新',5:'时间偏移标定',6:'空间外参标定',7:'多目标航迹优化',8:'概率安全约束',9:'成像质量评分',10:'运动模糊约束',11:'级联注意力特征',12:'对比学习损失',13:'域泛化约束',14:'伪标签生成',15:'主动学习采样',16:'Focal Loss',17:'类别有效样本权重',18:'复合训练损失',19:'缩放点积注意力',20:'卷积增强注意块',21:'前馈残差块',22:'多尺度加权融合',23:'点云到图像投影',24:'导线悬链线',25:'最小净空距离',26:'保守净空',27:'不确定度传播',28:'树冠生长预测',29:'树障风险指数',30:'边云任务卸载',31:'端到端时延',32:'系统可用度',33:'产学研收益分配',34:'直接经济效益',35:'净现值',36:'避免停电损失',37:'碳减排',38:'安全暴露降低'}
for k,v in formula_desc.items(): formula_rows.append([f'式（{k}）',v,'正文相应公式后已逐项解释变量；最终参数和阈值以试验、标准和企业规程为准'])
add_table(doc,['公式编号','模型用途','变量与使用边界'],formula_rows,[2.2,6.0,9.5],8.2)

add_heading(doc,'附录B 申报证据材料清单',1)
add_table(doc,['序号','证据材料','对应章节/创新点','状态与责任建议'],[
['1','立项批复、任务书、合作协议和会议纪要','开发历程、产学研合作','项目管理组核对时间、任务和单位'],
['2','算法数据集说明、训练日志、代码版本和对比试验','创新点三至五、技术水平','算法组形成可复现包'],
['3','标定、飞行、测距、可靠性和软件测试报告','创新点一、二、六、八','测试组按统一口径汇总'],
['4','CMA/CNAS第三方检测报告','客观评价、技术水平','选择资质覆盖相符机构'],
['5','科技查新和成果评价证书','查新、技术水平','逐条对应查新点和创新点'],
['6','专利、软著、论文、标准清单和证书','其他评价、成果支撑','知识产权组核对权属和贡献'],
['7','用户应用证明、任务日志、缺陷工单和复检结果','应用情况和效益','用户单位盖章并保留联系人'],
['8','合同、发票、财务账和专项审计','直接经济效益','财务组明确项目边界并剔除无关收入'],
['9','单位及个人贡献证明','完成单位与完成人排序','任务、代码、试验、成果和转化留痕'],
['10','网络安全、数据合规和飞行许可材料','平台和现场应用','安全/法务组归档']],[1.5,6.5,5.5,5.0])

add_heading(doc,'附录C 图件与Visio源文件对应表',1)
fig_rows=[]
for idx,title,*_ in fig_defs:
    fig_rows.append([f'图{idx}',title,fig_paths[idx][1],'正文已引用；VSDX中方框、文字和连线可编辑'])
add_table(doc,['图号','图名','Visio源文件','说明'],fig_rows,[1.5,6.5,6.5,3.5],8.4)

add_heading(doc,'参考文献与政策标准依据',1)
refs=[
'[1] 国家发展改革委、国家能源局. 配电网高质量发展及升级改造相关行动文件，2024.',
'[2] 国家能源局. 新型电力系统发展蓝皮书，2023.',
'[3] Faisal, M.A.A.; et al. Deep Learning in Automated Power Line Inspection: A Review. arXiv:2502.07826, 2025.',
'[4] Vieira e Silva, A.L.B.; et al. InsPLAD: A Dataset and Benchmark for Power Line Asset Inspection in UAV Images. arXiv:2311.01619, 2023.',
'[5] Yang, S.; et al. Power Line Aerial Image Restoration under Adverse Weather: Datasets and Baselines. arXiv:2409.04812, 2024.',
'[6] Ollero, A.; et al. AERIAL-CORE: AI-Powered Aerial Robots for Inspection and Maintenance of Electrical Power Infrastructures. arXiv:2401.02343, 2024.',
'[7] Rong, S.; et al. Advanced YOLO-based Real-time Power Line Detection for Vegetation Management. arXiv:2503.00044, 2025.',
'[8] Riss, V.; et al. Autonomous Inspection of Power Line Insulators with UAV on an Unmapped Transmission Tower. arXiv:2602.24011, 2026.',
'[9] Abdelfattah, R.; Wang, X.; Wang, S. TTPLA: An Aerial-Image Dataset for Detection and Segmentation of Transmission Towers and Power Lines. arXiv:2010.10032, 2020.',
'[10] Vieira-e-Silva, A.L.B.; et al. STN PLAD: A Dataset for Multi-Size Power Line Assets Detection in High-Resolution UAV Images. arXiv:2108.07944, 2021.',
'[11] Liu, Z.; et al. Swin Transformer: Hierarchical Vision Transformer Using Shifted Windows. ICCV, 2021.',
'[12] Chen, T.; et al. A Simple Framework for Contrastive Learning of Visual Representations. ICML, 2020.',
'[13] He, K.; et al. Momentum Contrast for Unsupervised Visual Representation Learning. CVPR, 2020.',
'[14] Carion, N.; et al. End-to-End Object Detection with Transformers. ECCV, 2020.',
'[15] Lin, T.-Y.; et al. Focal Loss for Dense Object Detection. ICCV, 2017.',
'[16] ISO/IEC 23894:2023. Information technology—Artificial intelligence—Guidance on risk management.',
'[17] GB/T 43555—2023. 智能服务 预测性维护 算法测评方法.',
'[18] GB/T 43370—2023. 民用无人机地理围栏数据技术规范.',
'[19] GB/T 43570—2023. 民用无人驾驶航空器系统身份识别 总体要求.',
'[20] 输配电线路运行、无人机作业、网络安全和数据管理相关现行国家、行业及企业标准。']
for ref in refs: add_p(doc,ref,indent=False,spacing=1.15,after=2)

# save
DOCX=OUT/'Optimized_Power_Line_UAV_Award_Application.docx'
doc.save(DOCX)
# verify native OMML and figures
with zipfile.ZipFile(DOCX) as z:
    xml=z.read('word/document.xml')
    assert xml.count(b'<m:oMath') >= 38, f'OMML count insufficient: {xml.count(b"<m:oMath")}'
    assert xml.count(b'<pic:pic') >= 15, f'Figure count insufficient: {xml.count(b"<pic:pic")}'
# validate vsdx package structure
for f in FIG_VSDX.glob('*.vsdx'):
    with zipfile.ZipFile(f) as z:
        names=set(z.namelist()); required={'[Content_Types].xml','visio/document.xml','visio/pages/pages.xml','visio/pages/page1.xml'}
        assert required.issubset(names), f'Invalid VSDX {f.name}'
# convert docx to PDF for QA
subprocess.run(['libreoffice','--headless','--convert-to','pdf','--outdir',str(OUT),str(DOCX)],check=False,stdout=subprocess.PIPE,stderr=subprocess.STDOUT)
# create readme and package
readme=OUT/'下载与使用说明.txt'
readme.write_text('''本压缩包包含：\n1. Optimized_Power_Line_UAV_Award_Application.docx：本轮优化后的完整Word申报材料。\n2. Visio_Editable_Figures：15幅VSDX可编辑源图。\n3. Figures_SVG：对应SVG矢量备份。\n4. Optimized_Power_Line_UAV_Award_Application.pdf：排版校验预览（如生成成功）。\n\n公式说明：正文含38个Word原生OMML公式，可在Microsoft Word中直接编辑。\n图件说明：VSDX文件采用独立矩形、文字和连线形状构建，可使用Microsoft Visio打开并编辑。\n证据边界：所有带“待补充/待核验”的指标必须在正式申报前由第三方报告、应用证明或专项审计核实。\n''',encoding='utf-8')
package=OUT/'Power_Line_UAV_Award_Optimized_Package.zip'
with zipfile.ZipFile(package,'w',zipfile.ZIP_DEFLATED) as z:
    z.write(DOCX,DOCX.name)
    pdf=OUT/'Optimized_Power_Line_UAV_Award_Application.pdf'
    if pdf.exists(): z.write(pdf,pdf.name)
    z.write(readme,readme.name)
    for f in sorted(FIG_VSDX.glob('*.vsdx')): z.write(f,f'Visio_Editable_Figures/{f.name}')
    for f in sorted(FIG_SVG.glob('*.svg')): z.write(f,f'Figures_SVG/{f.name}')
print('BUILD_OK',DOCX.stat().st_size,package.stat().st_size,len(list(FIG_VSDX.glob('*.vsdx'))))
